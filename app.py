import os
import io
import streamlit as st
from docx import Document
from PIL import Image
import pypdf
import google.generativeai as genai

# -------------------------------------------------------------
# 1. मोबाइल स्क्रीन एवं पेज कॉन्फ़िगरेशन
# -------------------------------------------------------------
st.set_page_config(
    page_title="खाद्य निरीक्षक डिजिटल सहायक",
    page_icon="🌾",
    layout="centered",
    initial_sidebar_state="collapsed"
)

st.markdown("""
    <style>
    .main { padding: 0.5rem; }
    .stButton>button { width: 100%; border-radius: 8px; height: 3em; font-weight: bold; font-size: 16px; }
    </style>
""", unsafe_allow_html=True)

# -------------------------------------------------------------
# 2. सुरक्षित API Key एवं मॉडल सेटअप
# -------------------------------------------------------------
raw_api_key = st.secrets.get("GEMINI_API_KEY", os.environ.get("GEMINI_API_KEY", ""))

with st.sidebar:
    st.header("⚙️ सेटिंग्स")
    if not raw_api_key:
        raw_api_key = st.text_input("Gemini API Key दर्ज करें:", type="password")

api_key = str(raw_api_key).strip().replace('"', '').replace("'", "")

if not api_key:
    st.warning("⚠️ कृपया जारी रखने के लिए अपनी Gemini API Key दर्ज करें।")
    st.stop()

# GenAI कॉन्फ़िगरेशन
genai.configure(api_key=api_key)

SYSTEM_PROMPT = """
आप छत्तीसगढ़ शासन के खाद्य, नागरिक आपूर्ति एवं उपभोक्ता संरक्षण विभाग के खाद्य निरीक्षक (Food Inspector) के लिए एक उच्चस्तरीय विशेषज्ञ एआई सहायक हैं।
आपकी जिम्मेदारियां:
1. छत्तीसगढ़ सार्वजनिक वितरण प्रणाली (नियंत्रण) आदेश, आवश्यक वस्तु अधिनियम 1955 (धारा 3/7), राष्ट्रीय खाद्य सुरक्षा अधिनियम (NFSA) 2013 एवं उपभोक्ता संरक्षण अधिनियम के तहत सटीक विधिक प्रारूप तैयार करना।
2. शुद्ध, मानक शासकीय प्रशासनिक हिन्दी में 'कारण बताओ सूचना पत्र', 'पंचनामा', 'जब्ती सूची' एवं 'जांच प्रतिवेदन' तैयार करना।
3. राशन दुकानों, राइस मिलों, धान खरीदी केंद्रों के स्टॉक व सीएमआर (CMR) का सटीक ऑडिट करना।
"""

# आपके अकाउंट के लिए सक्रिय मॉडल को स्वतः खोजना (Auto-Model Discovery)
@st.cache_resource
def get_working_model(key: str):
    try:
        available = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
        # प्राथमिकता क्रम
        for pref in ["gemini-2.0-flash", "gemini-1.5-flash", "gemini-1.5-pro", "gemini-pro"]:
            for m in available:
                if pref in m:
                    return genai.GenerativeModel(m, system_instruction=SYSTEM_PROMPT)
        if available:
            return genai.GenerativeModel(available[0], system_instruction=SYSTEM_PROMPT)
    except Exception:
        pass
    return genai.GenerativeModel("gemini-1.5-flash", system_instruction=SYSTEM_PROMPT)

model = get_working_model(api_key)

# Word (.docx) फाइल बनाने का फंक्शन
def create_docx(text_content: str, title: str) -> io.BytesIO:
    doc = Document()
    doc.add_heading(title, level=1)
    for line in text_content.split("\n"):
        doc.add_paragraph(line)
    docx_io = io.BytesIO()
    doc.save(docx_io)
    docx_io.seek(0)
    return docx_io

# केंद्रीय कॉलिंग फंक्शन
def generate_ai_response(contents) -> str:
    try:
        response = model.generate_content(contents)
        if response and response.text:
            return response.text
        return "⚠️ कोई आउटपुट प्राप्त नहीं हुआ।"
    except Exception as e:
        return f"⚠️ त्रुटि (Error): {str(e)}"

# -------------------------------------------------------------
# 3. मुख्य यूजर इंटरफेस एवं टैब्स
# -------------------------------------------------------------
st.title("🌾 खाद्य निरीक्षक डिजिटल सहायक")
st.caption("क्षेत्रीय निरीक्षण, स्टॉक सत्यापन, धान मॉनिटरिंग एवं विधिक प्रतिवेदन (छ.ग. खाद्य विभाग)")

tab1, tab2, tab3, tab4, tab5 = st.tabs([
    "📋 नोटिस/पंचनामा", 
    "⚖️ स्टॉक सत्यापन", 
    "🌾 धान व CMR", 
    "📑 जनशिकायत जांच", 
    "📖 नियम निर्देशिका"
])

# =============================================================
# टैब 1: नोटिस एवं पंचनामा प्रारूप
# =============================================================
with tab1:
    st.subheader("📋 निरीक्षण नोटिस व पंचनामा तैयार करें")
    
    doc_type = st.selectbox(
        "दस्तावेज़ का प्रकार:",
        ["कारण बताओ सूचना पत्र (Show Cause Notice)", "मौका पंचनामा (Panchanama)", "दुकान निरीक्षण टीप (Inspection Note)"]
    )
    
    entity_type = st.selectbox(
        "निरीक्षण की गई संस्था:",
        ["उचित मूल्य दुकान (FPS)", "राइस मिल (Rice Mill)", "धान उपार्जन केंद्र (Paddy Centre)", "केरोसिन/घरेलू गैस विक्रेता"]
    )
    
    col1, col2 = st.columns(2)
    with col1:
        entity_name = st.text_input("दुकान / समिति / मिल का नाम:", placeholder="उदा. जय मां शारदा प्रा. साख समिति")
        entity_id = st.text_input("संस्था / दुकान कोड:", placeholder="उदा. 432001005")
    with col2:
        location = st.text_input("स्थान / ग्राम / वार्ड:", placeholder="उदा. अंबिकापुर, जिला सरगुजा")
        sample_file = st.file_uploader("पुराना सैंपल प्रारूप (वैकल्पिक):", type=["txt", "pdf", "docx"])
    
    common_violations = st.multiselect(
        "पाई गई सामान्य अनियमितताएं:",
        [
            "स्टॉक पंजी एवं वितरण पंजी संधारित नहीं पाई गई",
            "भौतिक सत्यापन में खाद्यान्न स्टॉक कम/अधिक पाया गया",
            "दुकान के बाहर सूचना पटल / दर सूची प्रदर्शित नहीं थी",
            "ई-पॉस (ePDS) मशीन व भौतिक वितरण में भिन्नता",
            "तौल कांटा सत्यापित/मुद्रांकित नहीं पाया गया",
            "निर्धारित समयावधि में दुकान खुली नहीं पाई गई",
            "हितग्राहियों को पावती रसीद प्रदान नहीं की जा रही थी"
        ]
    )
    additional_notes = st.text_area("अतिरिक्त विवरण / गवाहों के नाम:", placeholder="उदा. मौके पर 8.50 क्विंटल चावल की कमी पाई गई। 2 गवाह उपस्थित थे...")
    
    if st.button("दस्तावेज़ प्रारूप तैयार करें", key="btn_draft"):
        if not entity_name:
            st.error("कृपया संस्था/दुकान का नाम दर्ज करें।")
        else:
            with st.spinner("प्रशासनिक प्रारूप तैयार किया जा रहा है..."):
                violations_text = "\n".join([f"- {v}" for v in common_violations])
                if additional_notes:
                    violations_text += f"\n- अतिरिक्त विवरण: {additional_notes}"
                
                template_instruction = ""
                if sample_file:
                    try:
                        sample_content = sample_file.getvalue().decode("utf-8", errors="ignore")
                        template_instruction = f"\nकृपया इस संलग्न सैंपल प्रारूप की भाषा व संरचना का पालन करें:\n{sample_content}\n"
                    except Exception:
                        pass
                
                prompt = f"""
                निम्नलिखित विवरण के आधार पर मानक प्रशासनिक हिन्दी में औपचारिक '{doc_type}' का प्रारूप तैयार करें:
                - संस्था का प्रकार: {entity_type}
                - संस्था का नाम: {entity_name}
                - संस्था कोड: {entity_id}
                - स्थान: {location}
                - पाई गई अनियमितताएं:
                {violations_text}
                {template_instruction}

                आवश्यक निर्देश:
                1. छत्तीसगढ़ सार्वजनिक वितरण प्रणाली (नियंत्रण) आदेश 2016 एवं आवश्यक वस्तु अधिनियम 1955 का विधिक संदर्भ दें।
                2. 3 दिवस में जवाब प्रस्तुत करने का स्पष्ट निर्देश रखें।
                3. हस्ताक्षर, पदमुद्रा एवं एसडीएम / सहायक खाद्य अधिकारी को प्रतिलिपि का स्थान रखें।
                """
                result = generate_ai_response(prompt)
                st.markdown("### तैयार प्रारूप:")
                st.text_area("कॉपी करें", value=result, height=350)
                
                docx_file = create_docx(result, doc_type)
                st.download_button(
                    label="📥 Word (.docx) फाइल डाउनलोड करें",
                    data=docx_file,
                    file_name=f"{entity_name}_{doc_type}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

# =============================================================
# टैब 2: राशन स्टॉक सत्यापन एवं अंतर गणना
# =============================================================
with tab2:
    st.subheader("⚖️ राशन स्टॉक सत्यापन कैलकुलेटर")
    
    commodity = st.selectbox("सामग्री (Commodity):", ["चावल - अंत्योदय", "चावल - प्राथमिकता", "शक्कर", "नमक", "केरोसिन (लीटर)"])
    
    col1, col2 = st.columns(2)
    with col1:
        opening_stock = st.number_input("प्रारंभिक स्टॉक (क्विंटल):", min_value=0.0, step=0.1)
        received_stock = st.number_input("माह में प्राप्त आवंटन (क्विंटल):", min_value=0.0, step=0.1)
    with col2:
        distributed_stock = st.number_input("ePDS मशीन अनुसार वितरण (क्विंटल):", min_value=0.0, step=0.1)
        physical_stock = st.number_input("मौके पर भौतिक स्टॉक (क्विंटल):", min_value=0.0, step=0.1)
        
    book_balance = opening_stock + received_stock - distributed_stock
    variance = physical_stock - book_balance
    
    st.metric("पंजी अनुसार शेष स्टॉक (Book Balance)", f"{book_balance:.2f} क्विंटल")
    if variance < 0:
        st.error(f"⚠️ स्टॉक में कमी (Shortage): {abs(variance):.2f} क्विंटल")
    elif variance > 0:
        st.warning(f"⚠️ अधिक स्टॉक (Surplus): {variance:.2f} क्विंटल")
    else:
        st.success("✅ स्टॉक पूर्णतः संतुलित")
        
    if st.button("स्टॉक जांच टीप तैयार करें", key="btn_audit"):
        with st.spinner("जांच टीप तैयार हो रही है..."):
            prompt = f"""
            सामग्री: {commodity}
            - प्रारंभिक स्टॉक: {opening_stock} क्विंटल
            - प्राप्त आवंटन: {received_stock} क्विंटल
            - कुल वितरण: {distributed_stock} क्विंटल
            - पंजी अनुसार देय शेष: {book_balance:.2f} क्विंटल
            - मौके पर पाया गया भौतिक स्टॉक: {physical_stock} क्विंटल
            - अंतर: {variance:.2f} क्विंटल ({'कमी पाई गई' if variance < 0 else 'अतिरिक्त पाया गया' if variance > 0 else 'समान'})

            इस आंकड़े के आधार पर खाद्य निरीक्षक हेतु एक संक्षिप्त जांच टीप तैयार करें जिसमें अंतर की स्थिति और छत्तीसगढ़ पीडीएस आदेश के तहत की जाने वाली आगामी कानूनी कार्रवाई का स्पष्ट उल्लेख हो।
            """
            audit_result = generate_ai_response(prompt)
            st.markdown(audit_result)

# =============================================================
# टैब 3: धान खरीदी एवं सीएमआर (CMR) मॉनिटरिंग
# =============================================================
with tab3:
    st.subheader("🌾 धान उठाव एवं कस्टम मिलिंग (CMR) ट्रैकर")
    
    col_a, col_b = st.columns(2)
    with col_a:
        paddy_lifted = st.number_input("राइस मिलर द्वारा धान उठाव (क्विंटल):", min_value=0.0, step=1.0)
        out_turn_ratio = st.number_input("चावल रिकवरी प्रतिशत (%):", value=67.0, step=0.5)
    with col_b:
        cmr_expected = (paddy_lifted * (out_turn_ratio / 100.0))
        cmr_deposited = st.number_input("जमा किया गया चावल / CMR (क्विंटल):", min_value=0.0, step=1.0)
        
    pending_cmr = cmr_expected - cmr_deposited
    st.metric("देय कुल चावल (Expected CMR)", f"{cmr_expected:.2f} क्विंटल")
    st.metric("जमा हेतु शेष चावल (Pending CMR)", f"{pending_cmr:.2f} क्विंटल", delta=f"{-pending_cmr:.2f} क्विंटल" if pending_cmr > 0 else "पूर्ण")
    
    if st.button("मिलर को स्मरण / चेतावनी पत्र बनाएं", key="btn_cmr"):
        with st.spinner("चेतावनी पत्र तैयार हो रहा है..."):
            prompt = f"""
            छत्तीसगढ़ के राइस मिलर को कस्टम मिलिंग चावल (CMR) जमा करने में विलंब हेतु औपचारिक स्मरण/चेतावनी पत्र तैयार करें:
            - धान उठाव मात्रा: {paddy_lifted} क्विंटल
            - देय चावल मात्रा ({out_turn_ratio}% दर से): {cmr_expected:.2f} क्विंटल
            - जमा चावल मात्रा: {cmr_deposited} क्विंटल
            - शेष जमा योग्य चावल: {pending_cmr:.2f} क्विंटल
            
            निर्देश: 3 दिवस में शेष चावल नागरिक आपूर्ति निगम/FCI में जमा करने, अन्यथा बैंक गारंटी राजसात करने और वसूली कार्रवाई की चेतावनी का उल्लेख करें।
            """
            cmr_notice = generate_ai_response(prompt)
            st.text_area("नोटिस का प्रारूप", value=cmr_notice, height=300)

# =============================================================
# टैब 4: जनशिकायत निराकरण एवं साक्ष्य दस्तावेज अपलोड
# =============================================================
with tab4:
    st.subheader("📑 दस्तावेज संलग्न कर विस्तृत जांच प्रतिवेदन तैयार करें")
    
    col1, col2 = st.columns(2)
    with col1:
        complaint_source = st.selectbox("शिकायत का स्रोत:", ["कलेक्टर जनदर्शन", "सीएम हेल्पलाइन / पीजी पोर्टल", "लोक सेवा गारंटी आवेदन", "सीधे प्राप्त शिकायत"])
        complaint_no = st.text_input("शिकायत संदर्भ क्रमांक:", placeholder="उदा. PG/2026/08/1124")
    with col2:
        complainant_name = st.text_input("शिकायतकर्ता का नाम:", placeholder="उदा. रमेश कुमार एवं अन्य ग्रामीण")
        target_entity = st.text_input("संबंधित उचित मूल्य दुकान:", placeholder="उदा. उचित मूल्य दुकान कोड 432002")
        
    uploaded_files = st.file_uploader(
        "📎 साक्ष्य दस्तावेज अपलोड करें (शिकायत पत्र, बयान, ई-पॉस पर्ची, पंजी की फोटो):",
        type=["png", "jpg", "jpeg", "pdf", "txt"],
        accept_multiple_files=True
    )
    
    investigation_notes = st.text_area("मौके पर जांच में पाए गए मुख्य बिंदु:", placeholder="उदा. मौके पर 5 हितग्राहियों के बयान लिए गए, स्टॉक में 4 क्विंटल कमी पाई गई...")
    
    if st.button("संपूर्ण जांच प्रतिवेदन (Inquiry Report) बनाएं", key="btn_inquiry"):
        if not complainant_name:
            st.error("कृपया शिकायतकर्ता का नाम दर्ज करें।")
        else:
            with st.spinner("साक्ष्यों एवं बयानों का विश्लेषण कर प्रतिवेदन तैयार किया जा रहा है..."):
                contents_payload = []
                
                prompt_inquiry = f"""
                छत्तीसगढ़ खाद्य विभाग के खाद्य निरीक्षक की ओर से एसडीएम/उच्चाधिकारी हेतु विस्तृत 'जांच प्रतिवेदन' तैयार करें:
                - शिकायत स्रोत: {complaint_source}
                - संदर्भ क्रमांक: {complaint_no}
                - शिकायतकर्ता: {complainant_name}
                - संबंधित दुकान: {target_entity}
                - निरीक्षक की टीप: {investigation_notes}
                
                निर्देश:
                1. संलग्न सभी दस्तावेजों (फोटो, पर्ची, बयान) का विश्लेषण करके तथ्य निकालें।
                2. मानक शासकीय प्रतिवेदन तैयार करें: विषय, संदर्भ, घटनाक्रम, गवाहों के बयान, विधिक समीक्षा, स्पष्ट निष्कर्ष, एवं संलग्नक सूची (Annexures)।
                """
                contents_payload.append(prompt_inquiry)
                
                # विभिन्न प्रकार की फाइलों (इमेज / PDF / टेक्स्ट) को सुरक्षित तरीके से पढ़ना
                if uploaded_files:
                    for f in uploaded_files:
                        fname = f.name.lower()
                        try:
                            if fname.endswith(('.png', '.jpg', '.jpeg')):
                                img = Image.open(f)
                                contents_payload.append(img)
                            elif fname.endswith('.pdf'):
                                pdf_reader = pypdf.PdfReader(io.BytesIO(f.getvalue()))
                                pdf_text = "\n".join([page.extract_text() or "" for page in pdf_reader.pages])
                                contents_payload.append(f"\n[संलग्न PDF फाइल ({f.name}) का टेक्स्ट]:\n{pdf_text}\n")
                            else:
                                text_data = f.getvalue().decode("utf-8", errors="ignore")
                                contents_payload.append(f"\n[संलग्न टेक्स्ट फाइल ({f.name})]:\n{text_data}\n")
                        except Exception as e:
                            st.warning(f"फाइल {f.name} पढ़ने में समस्या: {e}")
                
                inquiry_res = generate_ai_response(contents_payload)
                st.markdown("### 📄 तैयार जांच प्रतिवेदन:")
                st.text_area("प्रतिवेदन कॉपी करें", value=inquiry_res, height=450)
                
                inquiry_docx = create_docx(inquiry_res, "जांच प्रतिवेदन")
                st.download_button(
                    label="📥 जांच प्रतिवेदन Word (.docx) डाउनलोड करें",
                    data=inquiry_docx,
                    file_name=f"Inquiry_Report_{complaint_no}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

# =============================================================
# टैब 5: शासकीय नियम एवं परिपत्र निर्देशिका
# =============================================================
with tab5:
    st.subheader("📖 केंद्र व राज्य शासकीय नियम व परिपत्र खोज")
    st.caption("अधिनियम, नियंत्रण आदेश एवं विभागीय प्रक्रियाओं की जानकारी प्राप्त करें।")
    
    query = st.text_input("नियम / परिपत्र संबंधी प्रश्न दर्ज करें:", placeholder="उदा. छत्तीसगढ़ में राशन दुकान निलंबन की कानूनी प्रक्रिया क्या है?")
    if st.button("नियम व विधिक प्रावधान खोजें", key="btn_search_law"):
        if query:
            with st.spinner("शासकीय आदेशों एवं अधिनियमों का विश्लेषण किया जा रहा है..."):
                search_prompt = f"""
                छत्तीसगढ़ खाद्य निरीक्षक के कार्य के संदर्भ में निम्नलिखित प्रश्न का कानूनी व प्रक्रियात्मक उत्तर दें:
                प्रश्न: {query}
                
                आवश्यक वस्तु अधिनियम 1955, छत्तीसगढ़ पीडीएस नियंत्रण आदेश 2016, राष्ट्रीय खाद्य सुरक्षा अधिनियम 2013 एवं संबंधित विभागीय नियमों की धाराएं एवं क्रमांक सहित उत्तर दें।
                """
                search_res = generate_ai_response(search_prompt)
                st.markdown(search_res)
